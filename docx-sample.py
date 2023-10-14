from docx import Document
import secrets

def extract_docx(filepath: str):
    text = ""
    # テキストの抽出
    document = Document(filepath)
    for p in document.paragraphs:
        if p.text == '':
            continue
        text += p.text

    # 表の抽出
    for table in document.tables:
        # 行ごとにcellを取得
        for row in table.rows:
            for cell in row.cells:
                if cell.text == '':
                    continue
                text += cell.text
    
    return text

local_path = "./docs/your_file_name.docx"

with open(f'./results/docx-result-{secrets.token_urlsafe(6)}.txt', 'w') as file:
    file.write(extract_docx(local_path))
