from docx import Document

def extract_docx(filepath: str):
    text = ""
    # テキストの抽出
    document = Document(filepath)
    for p in document.paragraphs:
        if p.text == '':
            continue
        text += p.text

    # 表の抽出
    for table in document.tables:
        # 行ごとにcellを取得
        for row in table.rows:
            for cell in row.cells:
                if cell.text == '':
                    continue
                text += cell.text
    
    return text



print(extract_docx("./docs/sample.docx"))
